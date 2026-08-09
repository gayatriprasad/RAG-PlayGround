"""Bring-your-own-corpus parser — Skill 33.

Dispatches by file extension into a normalized List[Document] matching the
same schema produced by parsers/enterprise_bench.py, so uploaded documents
flow through the exact same chunk/embed/index pipeline as bench documents.
"""

from __future__ import annotations

import csv
import json
import logging
from pathlib import Path
from typing import List, Optional

from raglab.config import CorpusCfg, IngestCfg
from raglab.types import Document, Question

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".html", ".htm"}


class UploadParser:
    """Parses user-uploaded files into normalized Document objects."""

    def parse_upload(
        self, file_path: str, cfg: Optional[CorpusCfg] = None, ingest_cfg: Optional[IngestCfg] = None
    ) -> List[Document]:
        """Parse a single uploaded file into one or more Documents.

        Args:
            file_path: Path to the uploaded file on disk.
            cfg: CorpusCfg controlling source_type inference.
            ingest_cfg: IngestCfg controlling which PDF parser backend to use
                (Skill 51) — 'marker' and 'surya' route to those parsers
                (each gracefully falls back to pdfplumber if not installed).

        Returns:
            List of Document objects (normalized via DocumentNormalizer by the caller).
        """
        cfg = cfg or CorpusCfg()
        ingest_cfg = ingest_cfg or IngestCfg()
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Uploaded file not found: {path}")

        ext = path.suffix.lower()
        source_type = self._infer_source_type(path, cfg)

        if ext in (".txt", ".md"):
            docs = self._parse_text(path, source_type)
        elif ext == ".pdf":
            docs = self._parse_pdf_routed(path, cfg, ingest_cfg, source_type)
        elif ext == ".docx":
            docs = self._parse_docx(path, source_type)
        elif ext == ".csv":
            docs = self._parse_csv(path, source_type)
        elif ext in (".html", ".htm"):
            docs = self._parse_html(path, source_type)
        else:
            raise ValueError(
                f"Unsupported file type: '{ext}'. "
                f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        from raglab.parsers.normalizer import DocumentNormalizer

        return DocumentNormalizer().normalize(docs)

    def _infer_source_type(self, path: Path, cfg: CorpusCfg) -> str:
        if cfg.auto_detect_source_type and path.parent.name not in ("", "."):
            return path.parent.name
        return path.stem

    def _parse_text(self, path: Path, source_type: str) -> List[Document]:
        content = path.read_text(encoding="utf-8", errors="replace")
        return [self._make_doc(path, source_type, content)]

    def _parse_pdf_routed(
        self, path: Path, cfg: CorpusCfg, ingest_cfg: IngestCfg, source_type: str
    ) -> List[Document]:
        """Route PDF parsing to marker/surya/pdfplumber based on IngestCfg.parser
        (Skill 51). 'auto' prefers marker (best structured-Markdown output);
        both marker and surya fall back gracefully if not installed."""
        match ingest_cfg.parser:
            case "marker" | "auto":
                from raglab.parsers.marker_parser import MarkerParser

                return MarkerParser().parse(str(path), cfg)
            case "surya":
                from raglab.parsers.surya_parser import SuryaParser

                return SuryaParser().parse(str(path), cfg)
            case _:
                return self._parse_pdf(path, source_type)

    def _parse_pdf(self, path: Path, source_type: str) -> List[Document]:
        content = ""
        try:
            import pdfplumber

            with pdfplumber.open(str(path)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                content = "\n\n".join(pages).strip()
        except Exception as e:
            logger.warning(f"pdfplumber failed for {path.name}: {e}")

        if not content:
            logger.warning(
                f"No extractable text found in {path.name} (scanned/image-only PDFs "
                f"are not supported without OCR)."
            )
        return [self._make_doc(path, source_type, content)]

    def _parse_docx(self, path: Path, source_type: str) -> List[Document]:
        import docx

        d = docx.Document(str(path))
        content = "\n\n".join(p.text for p in d.paragraphs if p.text.strip())
        return [self._make_doc(path, source_type, content)]

    def _parse_csv(self, path: Path, source_type: str) -> List[Document]:
        docs = []
        with open(path, "r", encoding="utf-8", errors="replace", newline="") as f:
            reader = csv.DictReader(f)
            for i, row in enumerate(reader):
                content = "\n".join(f"{k}: {v}" for k, v in row.items() if v)
                if not content.strip():
                    continue
                docs.append(
                    Document(
                        id=f"{source_type}_{path.stem}_{i}",
                        content=content,
                        source_type=source_type,
                        metadata={"filename": path.name, "row": i},
                    )
                )
        return docs

    def _parse_html(self, path: Path, source_type: str) -> List[Document]:
        from bs4 import BeautifulSoup

        raw = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        content = soup.get_text(separator="\n").strip()
        return [self._make_doc(path, source_type, content)]

    def _make_doc(self, path: Path, source_type: str, content: str) -> Document:
        return Document(
            id=f"{source_type}_{path.stem}",
            content=content,
            source_type=source_type,
            metadata={
                "filename": path.name,
                "filepath": str(path),
                "size_bytes": path.stat().st_size,
                "uploaded": True,
            },
        )

    def parse_directory(self, dir_path: str, cfg: Optional[CorpusCfg] = None) -> List[Document]:
        """Parse every supported file under `dir_path` (recursive)."""
        cfg = cfg or CorpusCfg()
        base = Path(dir_path)
        if not base.exists():
            return []

        docs: List[Document] = []
        for file_path in sorted(base.rglob("*")):
            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                try:
                    docs.extend(self.parse_upload(str(file_path), cfg))
                except Exception as e:
                    logger.warning(f"Failed to parse uploaded file {file_path}: {e}")
        return docs


def parse_upload(file_path: str, cfg: Optional[CorpusCfg] = None) -> List[Document]:
    """Module-level convenience wrapper around UploadParser.parse_upload."""
    return UploadParser().parse_upload(file_path, cfg)


def load_user_questions(path: str) -> List[Question]:
    """Load user-supplied Q&A pairs from a JSONL or CSV file.

    JSONL rows: {"question": ..., "answer": ..., "source_type"?, "category"?}
    CSV columns: question,answer[,source_type,category]

    Malformed/incomplete rows are skipped and logged, not raised.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"User questions file not found: {p}")

    questions: List[Question] = []
    ext = p.suffix.lower()

    if ext == ".jsonl":
        with open(p, "r", encoding="utf-8") as f:
            for line_num, line in enumerate(f, 1):
                line = line.strip()
                if not line:
                    continue
                try:
                    row = json.loads(line)
                    q = _row_to_question(row, line_num)
                    if q:
                        questions.append(q)
                except json.JSONDecodeError as e:
                    logger.warning(f"Skipping malformed JSONL line {line_num} in {p.name}: {e}")
    elif ext == ".csv":
        with open(p, "r", encoding="utf-8", errors="replace", newline="") as f:
            reader = csv.DictReader(f)
            for line_num, row in enumerate(reader, 1):
                q = _row_to_question(row, line_num)
                if q:
                    questions.append(q)
    else:
        raise ValueError(f"Unsupported questions file type: '{ext}'. Use .jsonl or .csv")

    logger.info(f"Loaded {len(questions)} user questions from {p.name}")
    return questions


def _row_to_question(row: dict, line_num: int) -> Optional[Question]:
    question_text = (row.get("question") or row.get("text") or "").strip()
    answer_text = (row.get("answer") or row.get("ground_truth") or "").strip()

    if not question_text or not answer_text:
        logger.warning(f"Skipping row {line_num}: missing question or answer")
        return None

    return Question(
        id=row.get("id") or f"user_q_{line_num}",
        text=question_text,
        ground_truth=answer_text,
        source_type=row.get("source_type") or "user",
        category=row.get("category") or "factual",
    )
